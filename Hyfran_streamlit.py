import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import norm, lognorm, gamma, gumbel_r, genextreme, pearson3
from scipy.optimize import curve_fit
import docx
from docx.shared import Inches
import base64
import io


def log_pearson3(x, loc, scale, skew):
    return pearson3.pdf(x, skew, loc, scale)


def fit_distribution(distr, data):
    params = distr.fit(data)
    log_likelihood = np.sum(np.log(distr.pdf(data, *params)))
    k = len(params)
    n = len(data)

    aic = 2 * k - 2 * log_likelihood
    bic = k * np.log(n) - 2 * log_likelihood

    return aic, bic, params


distributions = {
    'Normal': norm,
    'Lognormal': lognorm,
    'Pearson Type 3': pearson3,
    'Gamma': gamma,
    'Gumbel': gumbel_r,
    'GEV': genextreme,
}


def generate_word_document(max_flow, aic_bic_params, best_aic_distr, best_bic_distr):
    # Create a Word document
    doc = docx.Document()
    doc.add_heading('Frequency Analysis of Maximum Flow in Rivers', 0)

    doc.add_heading('Best Distribution based on AIC and BIC:', level=1)
    doc.add_paragraph(
        f"Best distribution based on AIC: {best_aic_distr} (AIC: {aic_bic_params[best_aic_distr]['AIC']})")
    doc.add_paragraph(
        f"Best distribution based on BIC: {best_bic_distr} (BIC: {aic_bic_params[best_bic_distr]['BIC']})")

    doc.add_heading('AIC and BIC for each distribution:', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Distribution'
    hdr_cells[1].text = 'AIC'
    hdr_cells[2].text = 'BIC'

    for name, info in aic_bic_params.items():
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(info['AIC'])
        row_cells[2].text = str(info['BIC'])

    doc.add_heading('Individual Distribution Plots:', level=1)
    for name in aic_bic_params.keys():
        doc.add_picture(f'{name}_distribution.png',
                        width=docx.shared.Inches(6))

    return doc


def download_link(document, filename):
    with io.BytesIO() as buffer:
        document.save(buffer)
        buffer.seek(0)
        file = base64.b64encode(buffer.read()).decode('utf-8')
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{file}" download="{filename}">Download Word document</a>'


st.title('Analyse fréquentielle des débits de crues')

st.text("Cet outil utilise des librairies 'open-source' afin de déterminer la meilleure distribution pour votre échantillon.")

uploaded_file = st.file_uploader(
    "Importer un fichier CSV d'une seule colonne qui comprend l'ensemble de l'échantillon.", type="csv")

if uploaded_file is not None:
    data = pd.read_csv(uploaded_file, header=None, names=['flow'])
    max_flow = data['flow'].to_numpy()

    aic_bic_params = {}
    for name, distr in distributions.items():
        aic, bic, params = fit_distribution(distr, max_flow)
        aic_bic_params[name] = {'AIC': aic, 'BIC': bic, 'params': params}

    best_aic_distr = min(
        aic_bic_params, key=lambda x: aic_bic_params[x]['AIC'])
    best_bic_distr = min(
        aic_bic_params, key=lambda x: aic_bic_params[x]['BIC'])

    x = np.linspace(min(max_flow), max(max_flow), 1000)
    for name, info in aic_bic_params.items():
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.hist(max_flow, bins='auto', density=True,
                alpha=0.6, color='g', label='Histogram')

        params = info['params']
        if name == 'Log-Pearson Type 3':
            ax.plot(x, log_pearson3(x, *params), label=name)
        else:
            distr = distributions[name]
            ax.plot(x, distr.pdf(x, *params), label=name)

        ax.set_xlabel('Flow')
        ax.set_ylabel('Density')
        ax.legend(loc='best')
        plt.savefig(f'{name}_distribution.png', bbox_inches='tight')
        plt.close(fig)

    doc = generate_word_document(
        max_flow, aic_bic_params, best_aic_distr, best_bic_distr)
    st.markdown(download_link(doc, 'Frequency_Analysis.docx'),
                unsafe_allow_html=True)

else:
    st.info("Importer votre fichier CSV.")
